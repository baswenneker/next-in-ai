import os
import datetime
import docx
from docx import Document
from typing import List
import readtime
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from next_in_ai.OpenAISummarizer import OpenAISummarizer


class BatchSummarizer:
    def __init__(self, urls: List[str]):
        self.urls = urls

    def create_summary_document(self, debug=False, output_filename=None):
        print("📄 Creating the summary document")
        document = Document()
        if output_filename is None:
            date_str = datetime.datetime.now().strftime("%Y-%m-%d")
            output_filename = f"summaries-{date_str}.docx"

        for url in self.urls:
            print("\n\nSummarizing article: ", url)
            summarizer = OpenAISummarizer(url)
            summary = summarizer.summarize()

            if debug:
                print("--------------------------------------")
                if summary is None:
                    print(f"No summary available for {url}")
                else:
                    print("Summary: ", summary)
                print("--------------------------------------")

            if summary is None:
                document.add_paragraph(f"\n\n❌ No summary available for {url}.\n\n")
            else:
                first_line = summary.split("\n")[0]  # Get the first line of the summary

                # Add a heading
                document.add_heading(f"{first_line}", level=2)

                # Add the summary
                rest_of_summary = summary.replace(first_line, "").strip()
                document.add_paragraph(rest_of_summary)

                # Calculate reading time
                reading_time = readtime.of_text(summarizer.content)
                reading_time_minutes = reading_time.minutes

                # Add the "Lees het volledige artikel" text with hyperlink
                paragraph = document.add_paragraph()
                text = f"Lees het volledige artikel ({reading_time_minutes} minuten).\n\n{url}\n\n"
                self._add_hyperlink(paragraph, url, text, "0000FF", True)

            # document.add_page_break()

        print("✅ Writing summaries to file.")
        document.save(output_filename)

    def _add_hyperlink(self, paragraph, url, text, color, underline):
        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(
            qn("r:id"),
            r_id,
        )

        new_run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        if color:
            color_element = OxmlElement("w:color")
            color_element.set(qn("w:val"), color)
            run_properties.append(color_element)

        if underline:
            underline_element = OxmlElement("w:u")
            underline_element.set(qn("w:val"), "single")
            run_properties.append(underline_element)

        new_run.append(run_properties)
        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)
        hyperlink.append(new_run)

        paragraph._element.append(hyperlink)
        return hyperlink
